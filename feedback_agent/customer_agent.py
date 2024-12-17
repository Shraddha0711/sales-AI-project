import os
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.document_loaders import TextLoader
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import CharacterTextSplitter
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langgraph.graph import StateGraph, END
from typing_extensions import TypedDict
from typing import List, Dict, Any
from langchain.docstore.document import Document  
import docx
from typing_extensions import TypedDict
from dotenv import load_dotenv
load_dotenv()
import prompts

# Initialize LLMs and embeddings

# Customer service metrics
llm_empathy_score = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_clarity_and_conciseness = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_grammar_and_language = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_listening_score = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_problem_resolution_effectiveness = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_personalisation_index = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_conflict_management = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_response_time = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_customer_satisfiction_score = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_positive_sentiment_score = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_structure_and_flow = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_stuttering_words = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_feedback = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)

embeddings = OpenAIEmbeddings()

# extract text from docx file
def extract_text_from_docx(file_path):
    
    doc = docx.Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)


# Define the GraphState with necessary fields
class GraphState(TypedDict):
    transcript: str

    empathy_score :str
    clarity_and_conciseness : str
    grammar_and_language : str 
    listening_score :str
    problem_resolution_effectiveness : str
    personalisation_index : str
    conflict_management : str 
    response_time : str
    customer_satisfiction_score : str
    positive_sentiment_score : str
    structure_and_flow : str
    stuttering_words : str
    feedback : str

# Load your .doc file into the text
transcript_file_path = "book_summary.docx"
transcript_text = extract_text_from_docx(transcript_file_path)

# Load and split documents into chunks
documents = [Document(page_content=transcript_text)]  # Add more documents as necessary
text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
texts = text_splitter.split_documents(documents)

# Create the FAISS index
vectorstore = FAISS.from_documents(texts, embeddings)

# Node functions for each metric
def empathy_score(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.empathy_score_prompt(transcript))
    chain = prompt | llm_empathy_score | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    state["empathy_score"] = score
    return state

def clarity_and_conciseness(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.clarity_and_conciseness_prompt(transcript))
    chain = prompt | llm_clarity_and_conciseness | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    state["clarity_and_conciseness"] = score
    return state

def grammar_and_language(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.grammar_and_language_prompt(transcript))
    chain = prompt | llm_grammar_and_language | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    state["grammar_and_language"] = score
    return state

def listening_score(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.listening_score_prompt(transcript))
    chain = prompt | llm_listening_score | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    state["listening_score"] = score
    return state

def problem_resolution_effectiveness(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.problem_resolution_effectiveness_prompt(transcript))
    chain = prompt | llm_listening_score | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    state["listening_score"] = score
    return state

def personalisation_index(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.personalisation_index_prompt(transcript))
    chain = prompt | llm_personalisation_index | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    state["personalisation_index"] = score
    return state

def conflict_management(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.conflict_management_prompt(transcript))
    chain = prompt | llm_conflict_management | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    state["conflict_management"] = score
    return state

def response_time(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.response_time_prompt(transcript))
    chain = prompt | llm_response_time | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    state["response_time"] = score
    return state

def customer_satisfiction_score(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.customer_satisfiction_score_prompt(transcript))
    chain = prompt | llm_customer_satisfiction_score | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    state["customer_satisfiction_score"] = score
    return state

def positive_sentiment_score(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.positive_sentiment_score_prompt(transcript))
    chain = prompt | llm_positive_sentiment_score | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    state["positive_sentiment_score"] = score
    return state

def structure_and_flow(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.structure_and_flow_prompt(transcript))
    chain = prompt | llm_structure_and_flow | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    state["structure_and_flow"] = score
    return state

def stuttering_words(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.stuttering_words_prompt(transcript))
    chain = prompt | llm_stuttering_words | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    state["stuttering_words"] = score
    return state

def feedback(state):
    transcript = state["transcript"]
    relevant_docs = vectorstore.similarity_search(transcript, k=5)
    retrieved_docs = "\n".join([doc.page_content for doc in relevant_docs])
    combined_prompt = f"{prompts.feedback_prompt(transcript)}\nRetrieved Knowledge: {retrieved_docs}"
    chain = ChatPromptTemplate.from_template(combined_prompt) | llm_feedback | StrOutputParser()
    feedback_result = chain.invoke({"transcript": transcript})
    state["feedback"] = feedback_result
    return state

# Define the graph
workflow = StateGraph(GraphState)

# Add nodes to the graph
workflow.add_node("node_empathy_score", empathy_score)
workflow.add_node("node_clarity_and_conciseness", clarity_and_conciseness)
workflow.add_node("node_grammar_and_language", grammar_and_language)
workflow.add_node("node_listening_score", listening_score)
workflow.add_node("node_problem_resolution_effectiveness", problem_resolution_effectiveness)
workflow.add_node("node_personalisation_index", personalisation_index)
workflow.add_node("node_conflict_management", conflict_management)
workflow.add_node("node_response_time", response_time)
workflow.add_node("node_customer_satisfiction_score", customer_satisfiction_score)
workflow.add_node("node_positive_sentiment_score", positive_sentiment_score)
workflow.add_node("node_structure_and_flow", structure_and_flow)
workflow.add_node("node_stuttering_words", stuttering_words)
workflow.add_node("node_feedback", feedback)

# Define the edges
workflow.add_edge("node_empathy_score", "node_clarity_and_conciseness")
workflow.add_edge("node_clarity_and_conciseness", "node_grammar_and_language")
workflow.add_edge("node_grammar_and_language", "node_listening_score")
workflow.add_edge("node_listening_score", "node_problem_resolution_effectiveness")
workflow.add_edge("node_problem_resolution_effectiveness", "node_personalisation_index")
workflow.add_edge("node_personalisation_index", "node_conflict_management")
workflow.add_edge("node_conflict_management", "node_response_time")
workflow.add_edge("node_response_time", "node_customer_satisfiction_score")
workflow.add_edge("node_customer_satisfiction_score", "node_positive_sentiment_score")
workflow.add_edge("node_positive_sentiment_score", "node_structure_and_flow")
workflow.add_edge("node_structure_and_flow", "node_stuttering_words")
workflow.add_edge("node_stuttering_words", "node_feedback")
workflow.add_edge("node_feedback", END)

# Set the entry point
workflow.set_entry_point("node_empathy_score")

# Compile the graph
customer_graph = workflow.compile()
