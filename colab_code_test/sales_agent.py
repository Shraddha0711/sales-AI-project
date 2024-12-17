import os
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.document_loaders import TextLoader
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import CharacterTextSplitter
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langgraph.graph import StateGraph, END, START
from typing_extensions import TypedDict
from typing import List, Dict, Any,Annotated
from langchain.docstore.document import Document  
import docx
from typing_extensions import TypedDict
from langgraph.checkpoint.memory import MemorySaver
from operator import add
from dotenv import load_dotenv
load_dotenv()
import prompts

# Initialize LLMs and embeddings

# Sales metrics
llm_product_knowledge_score = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_persuasion_and_negotiation_skills = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_objection_handling = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_confidence_score = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_value_proposition = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_pitch_quality = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_call_to_action_effectiveness = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_questioning_technique = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_rapport_building = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_active_listening_skills = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_upselling_success_rate = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_engagement = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_stuttering_words = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)
llm_feedback = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.4)

embeddings = OpenAIEmbeddings()

# extract text from docx file
def extract_text_from_docx(file_path):
    
    doc = docx.Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)

# Load your .doc file into the text
transcript_file_path = "book_summary.docx"
transcript_text = extract_text_from_docx(transcript_file_path)

# Load and split documents into chunks
documents = [Document(page_content=transcript_text)]  # Add more documents as necessary
text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
texts = text_splitter.split_documents(documents)

# Create the FAISS index
vectorstore = FAISS.from_documents(texts, embeddings)


class GraphState(TypedDict):
    transcript: str  # Input remains immutable
    aggregate: Annotated[list[str], add]  # Aggregates results from parallel nodes

# Node functions for each metric
def product_knowledge_score(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.product_knowledge_score_prompt(transcript))
    chain = prompt | llm_product_knowledge_score | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"product_knowledge_score: {score}"
    return {"aggregate": [result]}

def persuasion_and_negotiation_skills(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.persuasion_and_negotiation_skills_prompt(transcript))
    chain = prompt | llm_persuasion_and_negotiation_skills | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"persuasion_and_negotiation_skills: {score}"
    return {"aggregate": [result]}

def objection_handling(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.objection_handling_prompt(transcript))
    chain = prompt | llm_objection_handling | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"objection_handling: {score}"
    return {"aggregate": [result]}

def confidence_score(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.confidence_score_prompt(transcript))
    chain = prompt | llm_confidence_score | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"confidence_score: {score}"
    return {"aggregate": [result]}

def value_proposition(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.value_proposition_prompt(transcript))
    chain = prompt | llm_value_proposition | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"value_proposition: {score}"
    return {"aggregate": [result]}

def pitch_quality(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.pitch_quality_prompt(transcript))
    chain = prompt | llm_pitch_quality | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"pitch_quality: {score}"
    return {"aggregate": [result]}

def call_to_action_effectiveness(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.call_to_action_effectiveness_prompt(transcript))
    chain = prompt | llm_call_to_action_effectiveness | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"call_to_action_effectiveness: {score}"
    return {"aggregate": [result]}

def questioning_technique(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.questioning_technique_prompt(transcript))
    chain = prompt | llm_questioning_technique | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"questioning_technique: {score}"
    return {"aggregate": [result]}

def rapport_building(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.rapport_building_prompt(transcript))
    chain = prompt | llm_rapport_building | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"rapport_building: {score}"
    return {"aggregate": [result]}

def active_listening_skills(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.active_listening_skills_prompt(transcript))
    chain = prompt | llm_active_listening_skills | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"active_listening_skills: {score}"
    return {"aggregate": [result]}

def upselling_success_rate(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.upselling_success_rate_prompt(transcript))
    chain = prompt | llm_upselling_success_rate | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"upselling_success_rate: {score}"
    return {"aggregate": [result]}

def engagement(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.engagement_prompt(transcript))
    chain = prompt | llm_engagement | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"engagement: {score}"
    return {"aggregate": [result]}

def stuttering_words(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.stuttering_words_prompt(transcript))
    chain = prompt | llm_stuttering_words | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"stuttering_words: {score}"
    return {"aggregate": [result]}

def feedback(state):
    transcript = state["transcript"]
    relevant_docs = vectorstore.similarity_search(transcript, k=5)
    retrieved_docs = "\n".join([doc.page_content for doc in relevant_docs])
    combined_prompt = f"{prompts.feedback_prompt(transcript)}\nRetrieved Knowledge: {retrieved_docs}"
    chain = ChatPromptTemplate.from_template(combined_prompt) | llm_feedback | StrOutputParser()
    feedback_result = chain.invoke({"transcript": transcript})
    result = f"feedback: {feedback_result}"
    return {"aggregate": [result]}

# Define the graph
workflow = StateGraph(GraphState)

# Add nodes to the graph
workflow.add_node("node_product_knowledge_score", product_knowledge_score)
workflow.add_node("node_persuasion_and_negotiation_skills", persuasion_and_negotiation_skills)
workflow.add_node("node_objection_handling", objection_handling)
workflow.add_node("node_confidence_score", confidence_score)
workflow.add_node("node_value_proposition", value_proposition)
workflow.add_node("node_pitch_quality", pitch_quality)
workflow.add_node("node_call_to_action_effectiveness", call_to_action_effectiveness)
workflow.add_node("node_questioning_technique", questioning_technique)
workflow.add_node("node_rapport_building", rapport_building)
workflow.add_node("node_active_listening_skills", active_listening_skills)
workflow.add_node("node_upselling_success_rate", upselling_success_rate)
workflow.add_node("node_engagement", engagement)
workflow.add_node("node_stuttering_words", stuttering_words)
workflow.add_node("node_feedback", feedback)

# Define the edges
workflow.add_edge(START, "node_product_knowledge_score")
workflow.add_edge(START, "node_persuasion_and_negotiation_skills")
workflow.add_edge(START, "node_objection_handling")
workflow.add_edge(START, "node_confidence_score")
workflow.add_edge(START, "node_value_proposition")
workflow.add_edge(START, "node_pitch_quality")
workflow.add_edge(START, "node_call_to_action_effectiveness")
workflow.add_edge(START, "node_questioning_technique")
workflow.add_edge(START, "node_rapport_building")
workflow.add_edge(START, "node_active_listening_skills")
workflow.add_edge(START, "node_upselling_success_rate")
workflow.add_edge(START, "node_engagement")
workflow.add_edge(START, "node_stuttering_words")
workflow.add_edge(START, "node_feedback")

# Feedback node aggregates results
workflow.add_edge(["node_product_knowledge_score",
                   "node_persuasion_and_negotiation_skills",
                   "node_objection_handling",
                   "node_confidence_score",
                   "node_value_proposition",
                   "node_pitch_quality",
                   "node_call_to_action_effectiveness",
                   "node_questioning_technique",
                   "node_rapport_building",
                   "node_active_listening_skills",
                   "node_upselling_success_rate",
                   "node_engagement",
                   "node_stuttering_words",
                   "node_feedback"
                   ],END)

memory = MemorySaver()

# Compile the graph
sales_graph = workflow.compile(checkpointer=memory)
