import os
import json
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import CharacterTextSplitter
from langchain_core.output_parsers import StrOutputParser
from langgraph.graph import StateGraph, START, END
from typing_extensions import TypedDict
from typing import Annotated, List
import prompts
import docx
from langchain.docstore.document import Document
from langchain.text_splitter import CharacterTextSplitter
from langgraph.graph import StateGraph, START, END
from langchain_community.vectorstores import FAISS
from langgraph.checkpoint.memory import MemorySaver
from operator import add
from dotenv import load_dotenv
load_dotenv()

# Customer service metrics
llm_empathy_score = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.3)
llm_clarity_and_conciseness = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.3)
llm_grammar_and_language = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.3)
llm_listening_score = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.3)
llm_problem_resolution_effectiveness = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.3)
llm_personalisation_index = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.3)
llm_conflict_management = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.3)
llm_response_time = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.3)
llm_customer_satisfiction_score = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.3)
llm_positive_sentiment_score = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.3)
llm_structure_and_flow = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.3)
llm_stuttering_words = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.3)
llm_active_listening_skills = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.3)
llm_rapport_building = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.3)
llm_engagement = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.3)
llm_feedback = ChatOpenAI(model = "gpt-3.5-turbo", temperature=0.3)

embeddings = OpenAIEmbeddings()

# extract text from docx file
def extract_text_from_docx(file_path):

    doc = docx.Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)

# Load your .doc file into the text
transcript_file_path = "book_summary.docx"
transcript_text = extract_text_from_docx(transcript_file_path)

# Load and split documents into chunks
documents = [Document(page_content=transcript_text)]  # Add more documents as necessary
text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
texts = text_splitter.split_documents(documents)

# Create the FAISS index
vectorstore = FAISS.from_documents(texts, embeddings)


class GraphState(TypedDict):
    transcript: str  # Input remains immutable
    aggregate: Annotated[list[str], add]  # Aggregates results from parallel nodes

def empathy_score(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.empathy_score_prompt(transcript))
    chain = prompt | llm_empathy_score | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"empathy_score: {score}"
    return {"aggregate": [result]}

def clarity_and_conciseness(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.clarity_and_conciseness_prompt(transcript))
    chain = prompt | llm_clarity_and_conciseness | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"clarity_and_conciseness: {score}"
    return {"aggregate": [result]}

def grammar_and_language(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.grammar_and_language_prompt(transcript))
    chain = prompt | llm_grammar_and_language | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"grammar_and_language: {score}"
    return {"aggregate": [result]}

def listening_score(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.listening_score_prompt(transcript))
    chain = prompt | llm_listening_score | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"listening_score: {score}"
    return {"aggregate": [result]}

def problem_resolution_effectiveness(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.problem_resolution_effectiveness_prompt(transcript))
    chain = prompt | llm_listening_score | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"problem_resolution_effectiveness: {score}"
    return {"aggregate": [result]}

def personalisation_index(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.personalisation_index_prompt(transcript))
    chain = prompt | llm_personalisation_index | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"personalisation_index: {score}"
    return {"aggregate": [result]}

def conflict_management(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.conflict_management_prompt(transcript))
    chain = prompt | llm_conflict_management | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"conflict_management: {score}"
    return {"aggregate": [result]}

def response_time(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.response_time_prompt(transcript))
    chain = prompt | llm_response_time | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"response_time: {score}"
    return {"aggregate": [result]}

def customer_satisfiction_score(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.customer_satisfiction_score_prompt(transcript))
    chain = prompt | llm_customer_satisfiction_score | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"customer_satisfiction_score: {score}"
    return {"aggregate": [result]}

def positive_sentiment_score(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.positive_sentiment_score_prompt(transcript))
    chain = prompt | llm_positive_sentiment_score | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"positive_sentiment_score: {score}"
    return {"aggregate": [result]}

def structure_and_flow(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.structure_and_flow_prompt(transcript))
    chain = prompt | llm_structure_and_flow | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"structure_and_flow: {score}"
    return {"aggregate": [result]}

def stuttering_words(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.stuttering_words_prompt(transcript))
    chain = prompt | llm_stuttering_words | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"stuttering_words: {score}"
    return {"aggregate": [result]}

def active_listening_skills(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.active_listening_skills_prompt(transcript))
    chain = prompt | llm_active_listening_skills | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"active_listening_skills: {score}"
    return {"aggregate": [result]}

def rapport_building(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.rapport_building_prompt(transcript))
    chain = prompt | llm_rapport_building | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"rapport_building: {score}"
    return {"aggregate": [result]}

def engagement(state):
    transcript = state["transcript"]
    prompt = ChatPromptTemplate.from_template(prompts.engagement_prompt(transcript))
    chain = prompt | llm_engagement | StrOutputParser()
    score = chain.invoke({"transcript": transcript})
    result = f"engagement: {score}"
    return {"aggregate": [result]}

def feedback(state):
    transcript = state["transcript"]
    relevant_docs = vectorstore.similarity_search(transcript, k=5)
    retrieved_docs = "\n".join([doc.page_content for doc in relevant_docs])
    combined_prompt = f"{prompts.feedback_prompt(transcript)}\nRetrieved Knowledge: {retrieved_docs}"
    chain = ChatPromptTemplate.from_template(combined_prompt) | llm_feedback | StrOutputParser()
    feedback_result = chain.invoke({"transcript": transcript})
    result = f"feedback: {feedback_result}"
    return {"aggregate": [result]}

workflow = StateGraph(GraphState)

# Add nodes
workflow.add_node("node_empathy_score", empathy_score)
workflow.add_node("node_clarity_and_conciseness", clarity_and_conciseness)
workflow.add_node("node_grammar_and_language", grammar_and_language)
workflow.add_node("node_listening_score", listening_score)
workflow.add_node("node_problem_resolution_effectiveness", problem_resolution_effectiveness)
workflow.add_node("node_personalisation_index", personalisation_index)
workflow.add_node("node_conflict_management", conflict_management)
workflow.add_node("node_response_time", response_time)
workflow.add_node("node_customer_satisfiction_score", customer_satisfiction_score)
workflow.add_node("node_positive_sentiment_score", positive_sentiment_score)
workflow.add_node("node_structure_and_flow", structure_and_flow)
workflow.add_node("node_stuttering_words", stuttering_words)
workflow.add_node("node_active_listening_skills", active_listening_skills)
workflow.add_node("node_rapport_building", rapport_building)
workflow.add_node("node_engagement", engagement)
workflow.add_node("node_feedback", feedback)

# Define parallel execution
workflow.add_edge(START, "node_empathy_score")
workflow.add_edge(START, "node_clarity_and_conciseness")
workflow.add_edge(START, "node_grammar_and_language")
workflow.add_edge(START, "node_listening_score")
workflow.add_edge(START, "node_problem_resolution_effectiveness")
workflow.add_edge(START, "node_personalisation_index")
workflow.add_edge(START, "node_conflict_management")
workflow.add_edge(START, "node_response_time")
workflow.add_edge(START, "node_customer_satisfiction_score")
workflow.add_edge(START, "node_positive_sentiment_score")
workflow.add_edge(START, "node_structure_and_flow")
workflow.add_edge(START, "node_stuttering_words")
workflow.add_edge(START, "node_active_listening_skills")
workflow.add_edge(START, "node_rapport_building")
workflow.add_edge(START, "node_engagement")
workflow.add_edge(START, "node_feedback")

# Feedback node aggregates results
workflow.add_edge(["node_empathy_score",
                   "node_clarity_and_conciseness",
                   "node_grammar_and_language",
                   "node_listening_score",
                   "node_problem_resolution_effectiveness",
                   "node_personalisation_index",
                   "node_conflict_management",
                   "node_response_time",
                   "node_customer_satisfiction_score",
                   "node_positive_sentiment_score",
                   "node_structure_and_flow",
                   "node_active_listening_skills",
                   "node_rapport_building",
                   "node_engagement",
                   "node_feedback"
                   ],END)
# workflow.add_edge("node_output", END)
memory = MemorySaver()

# Compile the graph
customer_graph = workflow.compile(checkpointer=memory)

